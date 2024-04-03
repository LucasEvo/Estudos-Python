# I cannot directly create a Word document as my Python environment does not support it.
# However, I can provide a Python code that uses the 'python-docx' library to create a Word document.
# The user can run this code on their local machine to generate the Word document.

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set the title
title = doc.add_heading('Currículo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Personal Information
doc.add_heading('Informações Pessoais', level=1)
doc.add_paragraph('Nome Completo: [Seu Nome Completo]')
doc.add_paragraph('Telefone: [Seu Telefone]')
doc.add_paragraph('E-mail: [Seu E-mail]')
doc.add_paragraph('LinkedIn: [Seu LinkedIn] (opcional)')

# Professional Objective
doc.add_heading('Objetivo Profissional', level=1)
doc.add_paragraph('[Breve descrição dos seus objetivos de carreira e o que você espera do estágio.]')

# Academic Formation
doc.add_heading('Formação Acadêmica', level=1)
doc.add_paragraph('Instituição: [Nome da Instituição]')
doc.add_paragraph('Curso: [Nome do Curso]')
doc.add_paragraph('Previsão de Conclusão: [Mês e Ano]')

# Professional Experience
doc.add_heading('Experiência Profissional', level=1)
doc.add_paragraph('[Descreva suas experiências anteriores, responsabilidades e conquistas.]')

# Technical Skills
doc.add_heading('Habilidades Técnicas', level=1)
doc.add_paragraph('Linguagens de Programação: [C, C++, Java, Python, etc.]')
doc.add_paragraph('Ferramentas e Tecnologias: [Git, Docker, etc.]')
doc.add_paragraph('Bancos de Dados: [SQL, NoSQL]')
doc.add_paragraph('Desenvolvimento Web: [HTML, CSS, JavaScript]')

# Academic or Personal Projects
doc.add_heading('Projetos Acadêmicos ou Pessoais', level=1)
doc.add_paragraph('[Descreva projetos relevantes ou contribuições para projetos open-source.]')

# Languages
doc.add_heading('Idiomas', level=1)
doc.add_paragraph('[Idiomas e níveis de proficiência.]')

# Extracurricular Activities
doc.add_heading('Atividades Extracurriculares', level=1)
doc.add_paragraph('[Participação em clubes de tecnologia, hackathons, etc.]')

# Courses and Certifications
doc.add_heading('Cursos e Certificações', level=1)
doc.add_paragraph('[Cursos online, workshops ou certificações relevantes para a área.]')

# Save the document
doc.save('Curriculo_Estagiario_Ciencia_da_Computacao.docx')

print("O modelo de currículo foi criado e salvo como 'Curriculo_Estagiario_Ciencia_da_Computacao.docx'")

